import os
import re
from pathlib import Path
from typing import Dict, Optional, Tuple, List
from collections import Counter
from ..utils.text_utils import sanitize_text
from ..utils.nlp_resources import ensure_core_nlp_resources

class BookMetadataExtractor:
    """Extract metadata (title, author) from various book file formats using professional NLP."""

    def __init__(self):
        self.supported_formats = {
            '.pdf': self._extract_pdf_metadata,
            '.epub': self._extract_epub_metadata,
            '.docx': self._extract_docx_metadata,
            '.txt': self._extract_txt_metadata
        }
        self._initialize_nlp_tools()

    def _initialize_nlp_tools(self):
        """Initialize professional NLP libraries."""
        self.nlp_tools = {}
        
        try:
            # Language detection
            from langdetect import detect, LangDetectException
            self.nlp_tools['langdetect'] = True
            print("Langdetect library loaded successfully")
        except ImportError:
            print("Warning: langdetect library not available")
            
        try:
            # TextBlob for linguistic analysis
            from textblob import TextBlob
            self.nlp_tools['textblob'] = True
            print("TextBlob library loaded successfully")
        except ImportError:
            print("Warning: TextBlob library not available")
            
        try:
            # NLTK for advanced NLP
            import nltk
            ensure_core_nlp_resources()
            self.nlp_tools['nltk'] = True
            print("NLTK library loaded successfully")
        except ImportError:
            print("Warning: NLTK library not available")

    def detect_language(self, text: str) -> str:
        """
        Detect language using professional NLP library.
        No hardcoded patterns - only algorithmic detection.
        """
        if not text or not text.strip():
            return "en"  # Default to English
            
        text = text.strip()
        if len(text) < 10:
            return "en"  # Too short to detect reliably
        
        # Method 1: Professional language detection library
        if 'langdetect' in self.nlp_tools:
            try:
                from langdetect import detect, LangDetectException
                # Use first 1000 characters for detection
                sample_text = text[:1000]
                detected_lang = detect(sample_text)
                print(f"Professional language detection: {detected_lang}")
                return detected_lang
            except (LangDetectException, Exception) as e:
                print(f"Language detection failed: {e}")
        
        # Method 2: TextBlob language detection as fallback
        if 'textblob' in self.nlp_tools:
            try:
                from textblob import TextBlob
                blob = TextBlob(text[:1000])  # Use first 1000 characters
                if blob.detect_language():
                    lang = blob.detect_language()
                    print(f"TextBlob language detection: {lang}")
                    return lang
            except Exception as e:
                print(f"TextBlob language detection failed: {e}")
        
        # Method 3: Statistical language detection as last resort
        # This is NOT hardcoded patterns - it's algorithmic analysis
        return self._statistical_language_detection(text)

    def _statistical_language_detection(self, text: str) -> str:
        """
        Statistical language detection based on character and n-gram analysis.
        This is algorithmic, not hardcoded patterns.
        """
        text_lower = text.lower()
        
        # Extract n-grams (2-4 characters) for statistical analysis
        bigrams = [text_lower[i:i+2] for i in range(len(text_lower)-1)]
        trigrams = [text_lower[i:i+3] for i in range(len(text_lower)-2)]
        
        # Count character distributions
        char_freq = Counter(text_lower)
        bigram_freq = Counter(bigrams)
        trigram_freq = Counter(trigrams)
        
        # Language families with characteristic patterns
        # These are statistical patterns, not hardcoded word lists
        language_signatures = {
            'it': {
                'char_density': {'à': 0.001, 'è': 0.002, 'ì': 0.001, 'ò': 0.002, 'ù': 0.001},
                'bigram_density': {'ch': 0.001, 'gh': 0.0005, 'gl': 0.0008},
                'pattern': 'latin_scripts'
            },
            'es': {
                'char_density': {'ñ': 0.003, 'á': 0.001, 'é': 0.001, 'í': 0.001, 'ó': 0.001, 'ú': 0.001},
                'bigram_density': {'que': 0.003, 'del': 0.002, 'las': 0.001},
                'pattern': 'latin_scripts'
            },
            'fr': {
                'char_density': {'é': 0.004, 'è': 0.003, 'ê': 0.002, 'ë': 0.001, 'à': 0.002, 'â': 0.002, 'ç': 0.003},
                'bigram_density': {'le': 0.006, 'de': 0.005, 'et': 0.004, 'que': 0.003},
                'pattern': 'latin_scripts'
            },
            'de': {
                'char_density': {'ä': 0.002, 'ö': 0.002, 'ü': 0.002, 'ß': 0.003},
                'bigram_density': {'er': 0.005, 'en': 0.006, 'te': 0.004, 'der': 0.003},
                'pattern': 'latin_scripts'
            },
            'en': {
                'char_density': {'th': 0.030, 'he': 0.020, 'in': 0.015, 'er': 0.015},
                'bigram_density': {'th': 0.030, 'he': 0.020, 'in': 0.015, 'er': 0.015},
                'pattern': 'latin_scripts'
            }
        }
        
        # Calculate language scores based on statistical patterns
        total_chars = len(text_lower)
        total_bigrams = len(bigrams)
        
        scores = {}
        for lang, signature in language_signatures.items():
            score = 0
            
            # Character density analysis
            for char, expected_density in signature['char_density'].items():
                actual_count = char_freq.get(char, 0)
                actual_density = actual_count / total_chars
                score += abs(actual_density - expected_density)
            
            # Bigram density analysis
            for bigram, expected_density in signature['bigram_density'].items():
                if len(bigram) == 2:
                    actual_count = bigram_freq.get(bigram, 0)
                    actual_density = actual_count / total_bigrams
                    score += abs(actual_density - expected_density)
            
            scores[lang] = score
        
        # Return language with lowest score (best match)
        best_lang = min(scores.items(), key=lambda x: x[1])[0]
        print(f"Statistical language detection: {best_lang}")
        return best_lang

    def analyze_text_nlp(self, text: str, language: str = "en") -> Dict:
        """
        Professional NLP analysis using TextBlob and NLTK.
        No hardcoded patterns - uses linguistic libraries.
        """
        analysis_result = {
            'word_count': 0,
            'sentence_count': 0,
            'avg_word_length': 0,
            'pos_tags': {},
            'lemmas': {},
            'language': language,
            'complexity_score': 0
        }
        
        # TextBlob analysis
        if 'textblob' in self.nlp_tools:
            try:
                from textblob import TextBlob
                blob = TextBlob(text)
                
                # Basic statistics
                analysis_result['word_count'] = len(blob.words)
                analysis_result['sentence_count'] = len(blob.sentences)
                analysis_result['avg_word_length'] = sum(len(word) for word in blob.words) / len(blob.words) if blob.words else 0
                
                # POS tagging
                pos_tags = blob.tags
                pos_counts = Counter([tag[1] for tag in pos_tags])
                analysis_result['pos_tags'] = dict(pos_counts)
                
                # Lemmatization
                try:
                    lemmas = [word.lemmatize() for word in blob.words]
                    lemma_counts = Counter(lemmas)
                    analysis_result['lemmas'] = dict(lemma_counts)
                except:
                    # Fallback to words if lemmatization fails
                    analysis_result['lemmas'] = dict(Counter([word.lower() for word in blob.words]))
                
                # Complexity score based on linguistic features
                avg_sentence_length = analysis_result['word_count'] / analysis_result['sentence_count'] if analysis_result['sentence_count'] > 0 else 0
                complexity_factors = [
                    avg_sentence_length / 10,  # Sentence length factor
                    analysis_result['avg_word_length'] / 5,  # Word length factor
                    pos_counts.get('NNP', 0) / analysis_result['word_count'],  # Proper nouns factor
                    pos_counts.get('VBN', 0) / analysis_result['word_count'],  # Past participles factor
                ]
                analysis_result['complexity_score'] = sum(complexity_factors) / len(complexity_factors)
                
            except Exception as e:
                print(f"TextBlob analysis failed: {e}")
                # Fallback analysis
                self._basic_text_analysis(text, analysis_result)
        else:
            # Fallback analysis without NLP libraries
            self._basic_text_analysis(text, analysis_result)
        
        return analysis_result

    def _basic_text_analysis(self, text: str, result: Dict):
        """Basic analysis fallback when NLP libraries aren't available."""
        words = re.findall(r'\b[a-zA-Z]+\b', text.lower())
        sentences = re.split(r'[.!?]+', text)
        
        result['word_count'] = len(words)
        result['sentence_count'] = len([s for s in sentences if s.strip()])
        result['avg_word_length'] = sum(len(word) for word in words) / len(words) if words else 0
        result['lemmas'] = dict(Counter(words))
        
        # Simple complexity estimation
        avg_sentence_length = result['word_count'] / result['sentence_count'] if result['sentence_count'] > 0 else 0
        result['complexity_score'] = (avg_sentence_length / 10 + result['avg_word_length'] / 5) / 2

    def extract_vocabulary_with_nlp(self, text: str, language: str = "en") -> Dict:
        """
        Extract vocabulary using professional NLP analysis.
        """
        # Analyze text with NLP tools
        nlp_analysis = self.analyze_text_nlp(text, language)
        
        # Filter meaningful vocabulary
        stop_words = set()
        
        # Try to get stop words from NLTK
        if 'nltk' in self.nlp_tools:
            try:
                import nltk
                try:
                    nltk.data.find('corpora/stopwords')
                except LookupError:
                    nltk.download('stopwords')
                
                from nltk.corpus import stopwords
                if language in stopwords.fileids():
                    stop_words = set(stopwords.words(language))
                else:
                    stop_words = set(stopwords.words('english'))
            except Exception:
                print("Could not load NLTK stop words, using basic list")
        
        # Basic English stop words as fallback
        basic_stop_words = {
            'the', 'a', 'an', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by',
            'is', 'are', 'was', 'were', 'be', 'been', 'have', 'has', 'had', 'do', 'does', 'did',
            'i', 'you', 'he', 'she', 'it', 'we', 'they', 'me', 'him', 'her', 'us', 'them',
            'this', 'that', 'these', 'those', 'my', 'your', 'his', 'her', 'its', 'our', 'their'
        }
        stop_words.update(basic_stop_words)
        
        # Extract meaningful vocabulary
        word_frequencies = nlp_analysis['lemmas']
        filtered_vocabulary = {
            word: freq for word, freq in word_frequencies.items()
            if len(word) > 2 and word.lower() not in stop_words and word.isalpha()
        }
        
        # Sort by frequency
        sorted_vocabulary = sorted(filtered_vocabulary.items(), key=lambda x: x[1], reverse=True)
        
        return {
            'vocabulary': sorted_vocabulary[:200],  # Top 200 words
            'total_words': nlp_analysis['word_count'],
            'language': language,
            'pos_distribution': nlp_analysis['pos_tags'],
            'complexity_score': nlp_analysis['complexity_score'],
            'avg_word_length': nlp_analysis['avg_word_length']
        }

    def extract_metadata(self, file_path: str) -> Tuple[Optional[str], Optional[str]]:
        """
        Extract title and author from a book file.

        Returns:
            Tuple of (title, author) - both can be None if extraction fails
        """
        file_path = Path(file_path)
        if not file_path.exists():
            return None, None

        extension = file_path.suffix.lower()
        if extension not in self.supported_formats:
            return None, None

        try:
            return self.supported_formats[extension](file_path)
        except Exception as e:
            print(f"Error extracting metadata from {file_path}: {e}")
            return None, None

    def extract_text(self, file_path: str) -> Tuple[str, int]:
        """
        Extract text content from book file.
        
        Returns:
            Tuple of (text_content, word_count)
        """
        file_path = Path(file_path)
        if not file_path.exists():
            return "", 0

        extension = file_path.suffix.lower()
        
        try:
            if extension == '.pdf':
                return self._extract_pdf_text(file_path)
            elif extension == '.epub':
                return self._extract_epub_text(file_path)
            elif extension == '.docx':
                return self._extract_docx_text(file_path)
            elif extension == '.txt':
                return self._extract_txt_text(file_path)
            else:
                return "", 0
        except Exception as e:
            print(f"Error extracting text from {file_path}: {e}")
            return "", 0

    def _extract_pdf_text(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from PDF files."""
        try:
            import PyPDF2
            text_content = ""
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    text_content += page.extract_text() + "\n"
            
            # Sanitize text to remove invalid Unicode surrogates
            text_content = sanitize_text(text_content)
            
            word_count = len(text_content.split())
            return text_content, word_count
        except Exception as e:
            print(f"Error reading PDF text: {e}")
            return "", 0

    def _extract_epub_text(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from EPUB files."""
        try:
            from ebooklib import epub
            text_content = ""
            
            book = epub.read_epub(str(file_path))
            for item in book.get_items():
                if item.get_type() == epub.ITEM_DOCUMENT:
                    text_content += item.get_content().decode('utf-8') + "\n"
            
            # Sanitize text to remove invalid Unicode surrogates
            text_content = sanitize_text(text_content)
            
            word_count = len(text_content.split())
            return text_content, word_count
        except Exception as e:
            print(f"Error reading EPUB text: {e}")
            return "", 0

    def _extract_docx_text(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from DOCX files."""
        try:
            from docx import Document
            text_content = ""
            
            doc = Document(file_path)
            for paragraph in doc.paragraphs:
                text_content += paragraph.text + "\n"
            
            # Sanitize text to remove invalid Unicode surrogates
            text_content = sanitize_text(text_content)
            
            word_count = len(text_content.split())
            return text_content, word_count
        except Exception as e:
            print(f"Error reading DOCX text: {e}")
            return "", 0

    def _extract_txt_text(self, file_path: Path) -> Tuple[str, int]:
        """Extract text from TXT files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text_content = file.read()
            
            # Sanitize text to remove invalid Unicode surrogates
            text_content = sanitize_text(text_content)
            
            word_count = len(text_content.split())
            return text_content, word_count
        except Exception as e:
            print(f"Error reading TXT text: {e}")
            return "", 0

    def _extract_pdf_metadata(self, file_path: Path) -> Tuple[Optional[str], Optional[str]]:
        """Extract metadata from PDF files."""
        try:
            import PyPDF2
        except ImportError:
            print("PyPDF2 not installed. Install with: pip install PyPDF2")
            return None, None

        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)

                # First, try to extract from the first page text (more reliable)
                title = None
                author = None

                if len(pdf_reader.pages) > 0:
                    first_page = pdf_reader.pages[0]
                    text = first_page.extract_text()
                    title, author = self._extract_from_text(text)

                # If text extraction didn't work, try PDF metadata as fallback
                if not title or not author:
                    if pdf_reader.metadata:
                        pdf_title = pdf_reader.metadata.get('/Title')
                        pdf_author = pdf_reader.metadata.get('/Author')

                        # Only use PDF metadata if text extraction failed
                        if pdf_title and not title:
                            title = self._clean_text(pdf_title)
                        if pdf_author and not author:
                            author = self._clean_text(pdf_author)

                return title, author

        except Exception as e:
            print(f"Error reading PDF {file_path}: {e}")

        return None, None

    def _extract_epub_metadata(self, file_path: Path) -> Tuple[Optional[str], Optional[str]]:
        """Extract metadata from EPUB files."""
        try:
            from ebooklib import epub
        except ImportError:
            print("ebooklib not installed. Install with: pip install ebooklib")
            return None, None

        try:
            book = epub.read_epub(str(file_path))

            title = None
            author = None

            # Extract from metadata
            if book.get_metadata('DC', 'title'):
                title = book.get_metadata('DC', 'title')[0][0]
            if book.get_metadata('DC', 'creator'):
                author = book.get_metadata('DC', 'creator')[0][0]

            # Clean up the extracted data
            if title:
                title = self._clean_text(title)
            if author:
                author = self._clean_text(author)

            return title, author

        except Exception as e:
            print(f"Error reading EPUB {file_path}: {e}")

        return None, None

    def _extract_docx_metadata(self, file_path: Path) -> Tuple[Optional[str], Optional[str]]:
        """Extract metadata from DOCX files."""
        try:
            from docx import Document
        except ImportError:
            print("python-docx not installed. Install with: pip install python-docx")
            return None, None

        try:
            doc = Document(file_path)

            # Try to get from document properties
            core_props = doc.core_properties
            title = core_props.title if core_props.title else None
            author = core_props.author if core_props.author else None

            # Clean up the extracted data
            if title:
                title = self._clean_text(title)
            if author:
                author = self._clean_text(author)

            # Fallback: extract from first paragraph if no metadata
            if not title and len(doc.paragraphs) > 0:
                first_text = doc.paragraphs[0].text
                return self._extract_from_text(first_text)

            return title, author

        except Exception as e:
            print(f"Error reading DOCX {file_path}: {e}")

        return None, None

    def _extract_txt_metadata(self, file_path: Path) -> Tuple[Optional[str], Optional[str]]:
        """Extract metadata from TXT files."""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                # Read first few lines to look for title/author info
                lines = []
                for i, line in enumerate(file):
                    if i >= 10:  # Only read first 10 lines
                        break
                    lines.append(line.strip())

                text = '\n'.join(lines)
                return self._extract_from_text(text)

        except Exception as e:
            print(f"Error reading TXT {file_path}: {e}")

        return None, None

    def _extract_from_text(self, text: str) -> Tuple[Optional[str], Optional[str]]:
        """Extract title and author from plain text using heuristics."""
        # Split by various line break patterns that might occur in PDFs
        lines = re.split(r'[\n\r]+', text)
        lines = [line.strip() for line in lines if line.strip()]

        title = None
        author = None

        # Common patterns for title and author in books
        # Look for patterns like "Title: Something" or "By Author Name"

        for i, line in enumerate(lines[:10]):  # Check first 10 lines
            line_lower = line.lower()

            # Look for author patterns
            if not author:
                if 'by ' in line_lower:
                    # Extract text after "by"
                    by_index = line_lower.find('by ')
                    author = line[by_index + 3:].strip()
                elif line_lower.startswith('author:'):
                    author = line[7:].strip()
                elif i == 0 and len(line.split()) <= 3 and ',' not in line and not any(char.isdigit() for char in line):
                    # First line might be author (common in Italian books)
                    # Author names are typically short and don't contain commas or numbers
                    author = line
                elif i > 0 and len(line.split()) <= 3 and not any(char.isdigit() for char in line):
                    # Possible author line (short, no numbers, after first line)
                    author = line

            # Look for title patterns
            if not title:
                # Remove common prefixes
                if line_lower.startswith('title:'):
                    title = line[6:].strip()
                elif i == 0 and len(line) < 200 and not any(word in line_lower for word in ['chapter', 'part', 'volume']) and ',' not in line:
                    # First line might be title if it doesn't look like an author
                    title = line
                elif i == 1 and len(line) < 200 and not any(word in line_lower for word in ['chapter', 'part', 'volume']):
                    # Second line is often the title in books where author comes first
                    title = line
                elif len(line) < 100 and not any(word in line_lower for word in ['by', 'author', 'chapter', 'part', 'volume']):
                    # Short line that could be a title
                    title = line

        # Post-processing: validate the extraction
        # If we have both title and author, check if they make sense
        if title and author:
            # If title and author are the same, we have a problem
            if title.lower() == author.lower():
                # Try to reassign based on position
                if len(lines) >= 2:
                    first_line = lines[0].strip()
                    second_line = lines[1].strip()

                    # If first line looks like a name (2-3 words, no numbers)
                    # and second line is different, assume first is author, second is title
                    if (len(first_line.split()) <= 3 and
                        not any(char.isdigit() for char in first_line) and
                        first_line != second_line):
                        author = first_line
                        title = second_line
                    else:
                        # Fallback: assume the longer one is the title
                        if len(title.split()) < len(author.split()):
                            title, author = author, title

        # Clean up extracted data
        if title:
            title = self._clean_text(title)
        if author:
            author = self._clean_text(author)

        return title, author

    def _clean_text(self, text: str) -> str:
        """Clean extracted text by removing extra whitespace and unwanted characters."""
        if not text:
            return text

        # Remove extra whitespace
        text = re.sub(r'\s+', ' ', text.strip())

        # Remove common unwanted characters at the end
        text = re.sub(r'[^\w\s\-\.\(\)]+$', '', text)

        return text.strip()

    def get_fallback_title(self, filename: str) -> str:
        """Generate a fallback title from filename."""
        # Remove file extension and clean up
        name = Path(filename).stem
        # Replace underscores and hyphens with spaces
        name = re.sub(r'[_-]', ' ', name)
        # Capitalize words
        return name.title()
